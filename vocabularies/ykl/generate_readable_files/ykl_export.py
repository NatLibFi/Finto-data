#!/usr/bin/env python3
from __future__ import annotations
import argparse, csv, json, re
from collections import defaultdict, deque
from pathlib import Path
from rdflib import Graph, RDF
from rdflib.namespace import SKOS
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

YKL_BASE='http://urn.fi/URN:NBN:fi:au:ykl:'

def natural_notation_key(s: str):
    if not s: return ((9, ''),)
    parts=re.split(r'(\d+)',s)
    return tuple((0, int(p)) if p.isdigit() else (1, p) for p in parts if p != '')

def local_id(uri:str):
    return uri.rsplit(':',1)[-1]

def set_outline_level(style, level:int):
    pPr=style.element.get_or_add_pPr(); outline=pPr.find(qn('w:outlineLvl'))
    if outline is None:
        outline=OxmlElement('w:outlineLvl'); pPr.append(outline)
    outline.set(qn('w:val'),str(min(level,8)))

def parse_graph(path:Path):
    g=Graph(); g.parse(path,format='turtle')
    concepts={}
    for s in set(g.subjects(RDF.type,SKOS.Concept)):
        uri=str(s)
        nots=[str(o) for o in g.objects(s,SKOS.notation)]
        labels={}
        for o in g.objects(s,SKOS.prefLabel): labels[o.language or '']=str(o)
        concepts[uri]={
            'uri':uri,'identifier':local_id(uri),'notation':nots[0] if nots else '',
            'label_fi':labels.get('fi',''),'label_sv':labels.get('sv',''),'label_en':labels.get('en',''),
            'broader_uris':[str(o) for o in g.objects(s,SKOS.broader)],
        }
    children=defaultdict(list); roots=[]
    anomalies={'multiple_broader':[],'missing_fi':[],'missing_notation':[],'broader_outside_dataset':[],'cycles':[],'duplicate_notation':[]}
    seen_not=defaultdict(list)
    for c in concepts.values():
        if not c['label_fi']: anomalies['missing_fi'].append(c['identifier'])
        if not c['notation']: anomalies['missing_notation'].append(c['identifier'])
        else: seen_not[c['notation']].append(c['identifier'])
        pins=[u for u in c['broader_uris'] if u in concepts]; pouts=[u for u in c['broader_uris'] if u not in concepts]
        if len(pins)>1: anomalies['multiple_broader'].append((c['identifier'],[concepts[u]['identifier'] for u in pins]))
        if pouts: anomalies['broader_outside_dataset'].append((c['identifier'],pouts))
        if pins:
            parent=sorted(pins,key=lambda u:(natural_notation_key(concepts[u]['notation']),concepts[u]['identifier']))[0]
            c['parent_uri']=parent; children[parent].append(uri:=c['uri'])
        else:
            c['parent_uri']=''; roots.append(c['uri'])
    for n,ids in seen_not.items():
        if len(ids)>1: anomalies['duplicate_notation'].append((n,sorted(ids)))
    def node_key(u):
        c=concepts[u]
        return (natural_notation_key(c['notation']),c['identifier'])
    for u in children: children[u].sort(key=node_key)
    # Explicit top-level order: main classes 0..9, fiction additions, form classes, then anything unexpected.
    preferred=[YKL_BASE+str(i) for i in range(10)]+[YKL_BASE+'fiktioluokka',YKL_BASE+'muotoluokka']
    ordered_roots=[u for u in preferred if u in roots]+sorted([u for u in roots if u not in preferred],key=node_key)
    state={}; stack=[]
    def dfs(u):
        state[u]=1; stack.append(u)
        for v in children.get(u,[]):
            if state.get(v,0)==0: dfs(v)
            elif state.get(v)==1:
                i=stack.index(v) if v in stack else 0
                anomalies['cycles'].append([concepts[x]['identifier'] for x in stack[i:]]+[concepts[v]['identifier']])
        stack.pop(); state[u]=2
    for u in ordered_roots:
        if state.get(u,0)==0: dfs(u)
    for u in concepts:
        if state.get(u,0)==0: dfs(u)
    depths={}; q=deque((u,0) for u in ordered_roots)
    while q:
        u,d=q.popleft()
        if u in depths and depths[u]<=d: continue
        depths[u]=d
        for v in children.get(u,[]): q.append((v,d+1))
    for u in concepts: depths.setdefault(u,0)
    rows=[]; visited=set()
    def walk(u):
        if u in visited:return
        visited.add(u); c=concepts[u]; p=concepts.get(c.get('parent_uri',''))
        rows.append({'notation':c['notation'],'identifier':c['identifier'],'label_fi':c['label_fi'],'label_sv':c['label_sv'],'label_en':c['label_en'],
                     'parent_notation':p['notation'] if p else '','parent_identifier':p['identifier'] if p else '',
                     'level':depths[u],'uri':c['uri']})
        for v in children.get(u,[]): walk(v)
    for u in ordered_roots: walk(u)
    for u in sorted(concepts,key=node_key): walk(u)
    return rows,anomalies

def display_line(r):
    return (r['notation']+' ' if r['notation'] else '')+r['label_fi']

def write_txt(rows,out):
    with open(out,'w',encoding='utf-8') as f:
        for r in rows: f.write('    '*r['level']+display_line(r)+'\n')

def write_md(rows,out):
    with open(out,'w',encoding='utf-8') as f:
        f.write('# YKL-luokitus\n\n')
        for r in rows:
            prefix=(f"**{r['notation']}** " if r['notation'] else '')
            f.write('  '*r['level']+'- '+prefix+r['label_fi']+'\n')

def write_csv(rows,out):
    fields=['notation','identifier','label_fi','label_sv','label_en','parent_notation','parent_identifier','level','uri']
    with open(out,'w',encoding='utf-8-sig',newline='') as f:
        w=csv.DictWriter(f,fieldnames=fields,delimiter=';'); w.writeheader(); w.writerows(rows)

def write_docx(rows,out):
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.8); sec.left_margin=Cm(2); sec.right_margin=Cm(2)
    t=doc.add_paragraph(style='Title'); t.add_run('YKL-luokitus')
    maxlevel=min(max((r['level'] for r in rows),default=0),8)
    for lvl in range(maxlevel+1):
        name=f'YKL taso {lvl+1}'; st=doc.styles[name] if name in doc.styles else doc.styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH)
        st.font.name='Aptos'; st.font.size=Pt(max(9,12-lvl*.45)); st.paragraph_format.left_indent=Cm(lvl*.65)
        st.paragraph_format.space_before=Pt(6 if lvl==0 else 1); st.paragraph_format.space_after=Pt(1); set_outline_level(st,lvl)
    for r in rows:
        p=doc.add_paragraph(style=f"YKL taso {min(r['level'],maxlevel)+1}")
        if r['notation']:
            rr=p.add_run(r['notation']); rr.bold=True; p.add_run(' ')
        else:
            rr=p.add_run(r['label_fi']); rr.bold=True
            continue
        p.add_run(r['label_fi'])
    doc.save(out)

def main():
    ap=argparse.ArgumentParser(description='Vie YKL SKOS/Turtle hierarkiaksi TXT-, Markdown-, CSV- ja DOCX-muotoon.')
    ap.add_argument('input',type=Path); ap.add_argument('-o','--output-dir',type=Path,default=Path('ykl_export'))
    a=ap.parse_args(); a.output_dir.mkdir(parents=True,exist_ok=True)
    rows,anom=parse_graph(a.input); base=a.output_dir/'ykl_hierarkia'
    write_txt(rows,base.with_suffix('.txt')); write_md(rows,base.with_suffix('.md')); write_csv(rows,base.with_suffix('.csv')); write_docx(rows,base.with_suffix('.docx'))
    json.dump(rows,open(a.output_dir/'ykl_master.json','w',encoding='utf-8'),ensure_ascii=False,indent=2)
    json.dump(anom,open(a.output_dir/'ykl_tarkistus.json','w',encoding='utf-8'),ensure_ascii=False,indent=2)
    print(f'Exported {len(rows)} concepts'); print(json.dumps({k:len(v) for k,v in anom.items()},ensure_ascii=False))
if __name__=='__main__':main()
